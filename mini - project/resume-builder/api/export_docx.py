from flask import request, send_file, jsonify
from docx import Document
import tempfile

def export_docx():
    data = request.get_json()

    try:
        doc = Document()

        # Full Name
        full_name = data.get("fullName", "John Doe")
        doc.add_heading(full_name, 0)

        # Contact Info
        email = data.get("email", "")
        phone = data.get("phone", "")
        doc.add_paragraph(f"📧 {email} | 📱 {phone}")

        # Summary
        summary = data.get("summary", "")
        if summary:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary)

        # Education
        doc.add_heading("Education", level=1)
        edu = data.get("education", {})
        doc.add_paragraph(f"{edu.get('degree', '')} — {edu.get('institution', '')}, {edu.get('year', '')}")

        # Experience
        doc.add_heading("Experience", level=1)
        exp = data.get("experience", {})
        doc.add_paragraph(f"{exp.get('jobTitle', '')} at {exp.get('company', '')} ({exp.get('years', '')})")
        doc.add_paragraph(exp.get("description", ""))

        # Skills
        skills = data.get("skills", [])
        if skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(skills))

        # Save as temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return send_file(tmp.name, as_attachment=True, download_name="resume.docx")

    except Exception as e:
        return jsonify({"error": str(e)}), 500
