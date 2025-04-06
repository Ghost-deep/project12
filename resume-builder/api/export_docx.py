from flask import request, send_file
from docx import Document
from io import BytesIO

def export_docx():
    data = request.get_json()
    doc = Document()

    doc.add_heading(data['fullName'], 0)
    doc.add_paragraph(f"Email: {data['email']} | Phone: {data['phone']}")
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(data['summary'])

    doc.add_heading("Education", level=1)
    edu = data['education']
    doc.add_paragraph(f"{edu['degree']} — {edu['institution']} ({edu['year']})")

    doc.add_heading("Experience", level=1)
    exp = data['experience']
    doc.add_paragraph(f"{exp['jobTitle']} at {exp['company']} ({exp['years']})")
    doc.add_paragraph(exp['description'])

    doc.add_heading("Skills", level=1)
    for skill in data['skills']:
        doc.add_paragraph(f"• {skill}", style='List Bullet')

    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return send_file(docx_file, as_attachment=True, download_name="resume.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
